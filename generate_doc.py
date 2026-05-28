#!/usr/bin/env python
import json
import sys
from src.agents.outline_architect import Outline
from src.agents.writer import generate_document
from src.models.schemas import VerifiedPaper

def main():
    # Load outline
    with open("outline.json", "r") as f:
        outline_data = json.load(f)
    outline = Outline(**outline_data)
    
    # Load verified papers
    with open("result.json", "r") as f:
        result_data = json.load(f)
    
    papers = []
    for vp_data in result_data.get("verified_papers", []):
        papers.append(VerifiedPaper(
            title=vp_data["title"],
            authors=vp_data.get("authors", []),
            year=vp_data.get("year", 0),
            doi=vp_data.get("doi"),
            url=vp_data.get("url", ""),
            citation_count=vp_data.get("citation_count", 0),
            venue=vp_data.get("venue", "Unknown"),
            abstract=vp_data.get("abstract_preview", ""),
            keep=True,
            reason=vp_data.get("reason", "")
        ))
    
    print(f"Generating document for: {outline.title}")
    print(f"Using {len(papers)} verified papers.")
    print("Writing each subsection with delays to avoid rate limits...")
    
    doc = generate_document(
        outline=outline,
        assignment_type="research paper",
        level=result_data["level"],
        papers=papers
    )
    
    # Save to Markdown
    with open("generated_paper.md", "w") as f:
        f.write(doc)
    print("Document saved to generated_paper.md")
    
    # Basic DOCX export
    try:
        from docx import Document
        docx = Document()
        docx.add_heading(outline.title, 0)
        lines = doc.split('\n')
        for line in lines:
            if line.startswith('# '):
                docx.add_heading(line[2:], 1)
            elif line.startswith('## '):
                docx.add_heading(line[3:], 2)
            elif line.startswith('### '):
                docx.add_heading(line[4:], 3)
            elif line.strip():
                docx.add_paragraph(line)
        docx.save("generated_paper.docx")
        print("DOCX saved to generated_paper.docx")
    except ImportError:
        print("python-docx not installed – skipping DOCX export")

if __name__ == "__main__":
    main()
