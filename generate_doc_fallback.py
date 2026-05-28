#!/usr/bin/env python
import json
from src.agents.outline_architect import Outline
from src.agents.writer_fallback import generate_document_fallback
from src.models.schemas import VerifiedPaper

with open("outline.json", "r") as f:
    outline_data = json.load(f)
outline = Outline(**outline_data)

with open("result.json", "r") as f:
    result_data = json.load(f)

papers = []
for vp_data in result_data.get("verified_papers", []):
    papers.append(VerifiedPaper(
        title=vp_data["title"],
        authors=vp_data.get("authors", []),
        year=vp_data.get("year", 0),
        doi=vp_data.get("doi"),
        url=vp_data.get("url", ""),
        citation_count=vp_data.get("citation_count", 0),
        venue=vp_data.get("venue", "Unknown"),
        abstract=vp_data.get("abstract_preview", ""),
        keep=True,
        reason=vp_data.get("reason", "")
    ))

doc = generate_document_fallback(outline, papers)

with open("generated_paper_fallback.md", "w") as f:
    f.write(doc)
print("Fallback document saved to generated_paper_fallback.md")

# Basic DOCX
try:
    from docx import Document
    docx = Document()
    docx.add_heading(outline.title, 0)
    for line in doc.split('\n'):
        if line.startswith('# '):
            docx.add_heading(line[2:], 1)
        elif line.startswith('## '):
            docx.add_heading(line[3:], 2)
        elif line.startswith('### '):
            docx.add_heading(line[4:], 3)
        elif line.strip():
            docx.add_paragraph(line)
    docx.save("generated_paper_fallback.docx")
    print("DOCX saved to generated_paper_fallback.docx")
except ImportError:
    pass
